"""Lecture du livrable produit, tel que le lecteur le verra (lot 4).

La passe de vérification ne relit pas les charges utiles des chapitres : elle
relit le **fichier livré**. C'est la leçon la plus chère du projet — le
markdown était propre, la barrière passait, et le document partait amputé parce
que quelque chose le refaisait après le contrôle (règle 3). Ce qui se vérifie
est ce qui se lit.

Le module ne juge rien. Il transforme un `.docx` en une matière comparable :
paragraphes hors tableaux, contenu des tableaux, et surtout **les grandeurs
chiffrées avec leur contexte**.
"""
from __future__ import annotations

import re
import statistics
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from core.numbers import (
    CURRENCY_ALTERNATION,
    MAGNITUDE_WORDS,
    NUMBER_BODY,
    SPACE_CLASS,
    parse_amount,
    parse_number,
)

#: Grandeur chiffrée : un nombre PORTANT une unité. Construit à partir des
#: briques de `core/numbers.py` plutôt que réécrit : la liste des devises n'a
#: qu'une seule source (règle 5). Le pourcentage y est ajouté ici, car
#: `core.numbers` ne le traite pas comme une unité monétaire.
_POURCENTAGE = r"%"
_GRANDEUR = re.compile(
    rf"(?<![A-Za-z\d])({NUMBER_BODY}){SPACE_CLASS}*"
    rf"({CURRENCY_ALTERNATION}|{MAGNITUDE_WORDS}|{_POURCENTAGE})",
    re.IGNORECASE,
)

#: Fenêtre de contexte reprise dans le motif d'anomalie. Un motif doit être
#: trouvable dans le document par le lecteur (règle 2) : sans le texte autour,
#: « 420 » est introuvable.
CONTEXTE = 60


@dataclass(frozen=True)
class Mesure:
    """Une grandeur chiffrée relevée dans le document livré."""

    valeur: float
    """Valeur ramenée à son unité de base : `1,25 M€` vaut 1 250 000."""
    unite: str
    """Unité telle qu'écrite dans le document."""
    texte: str
    """Le fragment exact, pour que le motif soit trouvable."""
    contexte: str
    """La phrase autour, pour situer le fragment."""
    dans_un_tableau: bool = False

    @property
    def est_un_pourcentage(self) -> bool:
        return self.unite.strip() == "%"

    @property
    def est_monetaire(self) -> bool:
        return not self.est_un_pourcentage and not self.est_une_magnitude_seule

    @property
    def est_une_magnitude_seule(self) -> bool:
        return self.unite.strip().lower() in ("million", "millions", "milliard", "milliards")


@dataclass
class DocumentLu:
    """Le livrable, ramené à ce qui est comparable."""

    chemin: Path
    paragraphes: list[str] = field(default_factory=list)
    """Prose hors tableaux."""
    cellules: list[str] = field(default_factory=list)
    """Contenu de toutes les cellules de tableau."""
    tableaux: int = 0
    tableaux_vides: int = 0
    images: int = 0
    mesures: list[Mesure] = field(default_factory=list)

    @property
    def texte_integral(self) -> str:
        return "\n".join([*self.paragraphes, *self.cellules])

    @property
    def mots(self) -> int:
        return len(self.texte_integral.split())

    @property
    def mots_en_tableaux(self) -> int:
        return sum(len(cellule.split()) for cellule in self.cellules)

    @property
    def part_en_tableaux(self) -> float:
        return self.mots_en_tableaux / self.mots if self.mots else 0.0

    @property
    def mediane_paragraphe(self) -> float:
        longueurs = [len(p.split()) for p in self.paragraphes if p.strip()]
        return statistics.median(longueurs) if longueurs else 0.0

    @property
    def part_paragraphes_longs(self) -> float:
        longueurs = [len(p.split()) for p in self.paragraphes if p.strip()]
        if not longueurs:
            return 0.0
        return sum(1 for n in longueurs if n > 60) / len(longueurs)


def _contexte(texte: str, debut: int, fin: int) -> str:
    extrait = texte[max(debut - CONTEXTE, 0) : fin + CONTEXTE]
    return " ".join(extrait.split())


def mesures_dans(texte: str, *, dans_un_tableau: bool = False) -> list[Mesure]:
    """Toutes les grandeurs chiffrées d'un texte, avec leur contexte.

    Ne relève QUE les nombres portant une unité. C'est une restriction
    délibérée, et il faut la connaître pour savoir ce que la passe ne voit
    pas : « trois portes d'entrée », « 0-30 j », « chapitre 12 » ne sont pas
    des affirmations de marché, et les traiter comme telles produirait des
    motifs faux — pires qu'absents (règle 2).
    """
    relevees: list[Mesure] = []
    for correspondance in _GRANDEUR.finditer(texte):
        brut, unite = correspondance.group(1), correspondance.group(2)
        if unite.strip() == "%":
            valeur = parse_number(brut)
        else:
            valeur = parse_amount(brut, unite)
        if valeur is None:
            continue
        relevees.append(
            Mesure(
                valeur=valeur,
                unite=unite,
                texte=correspondance.group(0).strip(),
                contexte=_contexte(texte, *correspondance.span()),
                dans_un_tableau=dans_un_tableau,
            )
        )
    return relevees


def lire_livrable(chemin: Path) -> DocumentLu:
    """Ouvre le `.docx` livré et en extrait la matière vérifiable.

    Lève si le fichier est absent ou illisible : un contrôle qui n'a rien à
    comparer est un échec, jamais un succès (règle 1).
    """
    from docx import Document  # noqa: PLC0415

    chemin = Path(chemin)
    if not chemin.is_file():
        msg = f"Livrable introuvable : {chemin}. Rien à vérifier."
        raise FileNotFoundError(msg)

    document = Document(str(chemin))
    lu = DocumentLu(chemin=chemin)

    for paragraphe in document.paragraphs:
        texte = paragraphe.text.strip()
        if texte:
            lu.paragraphes.append(texte)

    for table in document.tables:
        lu.tableaux += 1
        contenu_table: list[str] = []
        for ligne in table.rows:
            for cellule in ligne.cells:
                texte = cellule.text.strip()
                if texte:
                    contenu_table.append(texte)
        if not contenu_table:
            # Un tableau sans une seule cellule remplie est le symptôme exact
            # de la perte de lignes déjà constatée sur ce projet.
            lu.tableaux_vides += 1
        lu.cellules.extend(contenu_table)

    with zipfile.ZipFile(chemin) as archive:
        lu.images = sum(
            1 for nom in archive.namelist() if nom.startswith("word/media/")
        )

    for prose in lu.paragraphes:
        lu.mesures.extend(mesures_dans(prose))
    for contenu in lu.cellules:
        lu.mesures.extend(mesures_dans(contenu, dans_un_tableau=True))

    return lu


#: Balises dont le texte est de la PROSE. Les titres en font partie : le
#: contrôle d'intégrité cherche « CHAPITRE 07 » dans le texte intégral, et ce
#: marqueur vit dans un `<h…>`.
_BALISES_DE_PROSE = ("p", "li", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote")


def lire_livrable_html(html: str, *, chemin: Path | None = None) -> DocumentLu:
    """Même matière vérifiable, autre porte d'entrée : le document HTML.

    Le business plan et la stratégie tournent sur le moteur hérité : ils ne
    produisent ni socle ni `.docx`, donc `lire_livrable` — qui ouvre un `.docx`
    — ne peut rien lire chez eux. Deux des six contrôles du lot 4 ne dépendent
    pourtant pas du socle (intégrité et densité), et ce sont exactement ceux qui
    ont attrapé les désastres historiques du projet : tableaux vidés de leurs
    lignes, chapitres absents, document sans prose. Les priver de porte
    d'entrée revenait à ne pas les exécuter du tout.

    Deux précautions, toutes deux payées comptant sur ce dépôt :

    - **Les cellules sont relevées AVANT de détacher le tableau.** `decompose()`
      détruit l'élément ET ses enfants : c'est ce qui a envoyé un compte de
      résultat vide à un client (règle 3). On extrait, puis on détache — et avec
      `extract()`, qui détache sans détruire.
    - **On ne mesure pas notre propre balisage.** Le texte est pris balise par
      balise, jamais par un `get_text()` global : sans quoi `px`, `padding` et
      `cccccc` des styles en ligne entreraient dans la prose et fausseraient la
      densité — défaut réellement mesuré ici (corollaire de la règle 9).
    """
    from bs4 import BeautifulSoup  # noqa: PLC0415

    soupe = BeautifulSoup(html, "html.parser")
    for parasite in soupe(["script", "style", "head"]):
        parasite.decompose()

    lu = DocumentLu(chemin=chemin or Path("(document HTML)"))
    lu.images = len(soupe.find_all("img"))

    # Tableaux les plus extérieurs seulement : un tableau imbriqué verrait
    # sinon ses cellules comptées deux fois, et gonflerait la part en tableaux.
    tableaux = [t for t in soupe.find_all("table") if t.find_parent("table") is None]
    for tableau in tableaux:
        lu.tableaux += 1
        contenu = [
            cellule.get_text(" ", strip=True)
            for cellule in tableau.find_all(["td", "th"])
        ]
        contenu = [texte for texte in contenu if texte]
        if not contenu:
            lu.tableaux_vides += 1
        lu.cellules.extend(contenu)
        tableau.extract()

    for bloc in soupe.find_all(_BALISES_DE_PROSE):
        texte = bloc.get_text(" ", strip=True)
        if texte:
            lu.paragraphes.append(texte)

    for prose in lu.paragraphes:
        lu.mesures.extend(mesures_dans(prose))
    for contenu_cellule in lu.cellules:
        lu.mesures.extend(mesures_dans(contenu_cellule, dans_un_tableau=True))

    return lu
