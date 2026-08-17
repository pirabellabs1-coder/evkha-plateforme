"""Bibliothèque des composants du livrable (lot 0).

Chaque composant est une fonction testable indépendamment. Le générateur ne
fait qu'appeler ces fonctions dans l'ordre décrit par le JSON de l'étude.

Toutes les formes sont celles relevées dans `references/joalie_2026.docx` :

- bandeau de chapitre : tableau 1×1, fond prune               — 22 occurrences
- encadré étiqueté    : tableau 1×2, libellé prune + corps crème — 39
- grille de chiffres  : tableau 1×3, crème / rose / crème      —  4
- tableau de données  : en-tête prune, corps crème
- graphique           : image PNG matplotlib, pleine largeur   — 10
- couverture          : rectangle pleine page ancré derrière le texte

Pièges du lot 0, traités ici et nulle part ailleurs : largeurs posées en DXA
sur le tableau ET sur chaque cellule ; remplissage `clear` (`solid` rend en
noir) ; aucune puce en dur.
"""
from __future__ import annotations

import io
import re
from collections.abc import Sequence
from typing import Any

from docx.document import Document as DocumentWord
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.run import Run

from .gabarit import (
    STYLE_BANDEAU,
    STYLE_CHIFFRE_LIBELLE,
    STYLE_CHIFFRE_VALEUR,
    STYLE_CORPS,
    STYLE_ENCADRE_CORPS,
    STYLE_ENCADRE_TITRE,
    STYLE_LEGENDE,
    STYLE_SECTION,
    STYLE_SOURCE,
    STYLE_SOUS_TITRE,
    STYLE_TABLEAU_CELLULE,
    STYLE_TABLEAU_ENTETE,
    STYLE_TITRE_DOCUMENT,
)
from .palette import Palette, texte_lisible_sur, vers_rvb

#: Police de chaque style. La référence pose la police sur les RUNS et pas
#: seulement sur les styles : sans cela, Word peut substituer une fonte.
_POLICE_PAR_STYLE = {
    STYLE_TITRE_DOCUMENT: "Georgia",
    STYLE_BANDEAU: "Georgia",
    STYLE_CHIFFRE_VALEUR: "Georgia",
}
POLICE_CORPS = "Aptos"


def _poser_police(run: Run, style: str) -> None:
    run.font.name = _POLICE_PAR_STYLE.get(style, POLICE_CORPS)

#: Largeur utile en twips : A4 (11906) moins les marges (2 × 1134).
LARGEUR_UTILE_DXA = 9638
LARGEUR_UTILE_EMU = Emu(int(LARGEUR_UTILE_DXA / 1440 * 914400))

#: A4 en points, pour les formes pleine page de la couverture.
PAGE_LARGEUR_PT = 595.3
PAGE_HAUTEUR_PT = 841.9


# ── Utilitaires bas niveau ───────────────────────────────────────────────────


def _rgb(couleur: str) -> RGBColor:
    r, v, b = vers_rvb(couleur)
    return RGBColor(r, v, b)


def fond_cellule(cellule: _Cell, couleur: str) -> None:
    """Aplat de couleur. `clear` et jamais `solid` : `solid` rend en noir."""
    proprietes = cellule._tc.get_or_add_tcPr()
    proprietes.append(
        proprietes.makeelement(
            qn("w:shd"),
            {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): couleur.lstrip("#"),
            },
        )
    )


def _bordures(table: Table, couleur: str | None, epaisseur: int = 4) -> None:
    proprietes = table._tbl.tblPr
    bordures = proprietes.makeelement(qn("w:tblBorders"), {})
    for cote in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if couleur is None:
            attributs = {qn("w:val"): "none", qn("w:sz"): "0"}
        else:
            attributs = {
                qn("w:val"): "single",
                qn("w:sz"): str(epaisseur),
                qn("w:color"): couleur.lstrip("#"),
            }
        bordures.append(bordures.makeelement(qn(f"w:{cote}"), attributs))
    proprietes.append(bordures)


def marges_cellules(table: Table, haut: int = 80, cote: int = 120) -> None:
    """Marges intérieures, en twips."""
    proprietes = table._tbl.tblPr
    marges = proprietes.makeelement(qn("w:tblCellMar"), {})
    for nom, valeur in (("top", haut), ("bottom", haut), ("left", cote), ("right", cote)):
        marges.append(
            marges.makeelement(
                qn(f"w:{nom}"), {qn("w:w"): str(valeur), qn("w:type"): "dxa"}
            )
        )
    proprietes.append(marges)


def _largeur_table(table: Table, dxa: int) -> None:
    """Largeur en DXA sur le TABLEAU. Les pourcentages cassent le rendu."""
    proprietes = table._tbl.tblPr
    for existant in proprietes.findall(qn("w:tblW")):
        proprietes.remove(existant)
    proprietes.append(
        proprietes.makeelement(qn("w:tblW"), {qn("w:w"): str(dxa), qn("w:type"): "dxa"})
    )
    table.autofit = False


def _largeur_cellule(cellule: _Cell, dxa: int) -> None:
    """Largeur en DXA sur la CELLULE. Indispensable en plus de celle du tableau."""
    proprietes = cellule._tc.get_or_add_tcPr()
    for existant in proprietes.findall(qn("w:tcW")):
        proprietes.remove(existant)
    proprietes.append(
        proprietes.makeelement(qn("w:tcW"), {qn("w:w"): str(dxa), qn("w:type"): "dxa"})
    )


def ne_pas_fractionner(table: Table, *, entete_repetee: bool = False) -> None:
    """Interdit à une LIGNE de se couper entre deux pages.

    Word fractionne une ligne par défaut : une cellule de trois lignes de texte
    commencée en bas de page laisse sa première ligne seule, le reste passant à
    la page suivante. Sur un document de tableaux, c'est la principale source
    de pages à moitié vides — mesuré ici : 24 pages sous trente mots sur 67.

    `entete_repetee` redonne l'en-tête en haut de chaque page quand un tableau
    long finit quand même par s'étendre. Sans lui, le lecteur voit des colonnes
    sans savoir ce qu'elles portent.
    """
    for index, ligne in enumerate(table.rows):
        proprietes = ligne._tr.get_or_add_trPr()
        proprietes.append(OxmlElement("w:cantSplit"))
        if entete_repetee and index == 0:
            proprietes.append(OxmlElement("w:tblHeader"))


def garder_avec_la_suite(paragraphe: Any) -> None:
    """Attache un paragraphe à ce qui le suit.

    Un sous-titre seul en bas de page, son contenu à la page suivante : c'est
    l'autre moitié des pages creuses. `keep_with_next` déplace le titre avec
    son bloc au lieu de laisser la coupure les séparer.
    """
    paragraphe.paragraph_format.keep_with_next = True
    paragraphe.paragraph_format.keep_together = True


def _table(
    document: DocumentWord, lignes: int, colonnes: int, largeurs: Sequence[int]
) -> Table:
    table = document.add_table(rows=lignes, cols=colonnes)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _largeur_table(table, sum(largeurs))
    _bordures(table, None)
    marges_cellules(table)
    for ligne in table.rows:
        for index, cellule in enumerate(ligne.cells):
            _largeur_cellule(cellule, largeurs[index])
    # Tous les composants passent par ici — encadrés, grilles de chiffres,
    # tableaux de données, matrices. Poser la règle à cet endroit évite de
    # l'oublier sur l'un d'eux (règle 4 : viser la classe, pas les cas connus).
    ne_pas_fractionner(table)
    return table  # type: ignore[no-any-return]


def _ecrire(
    cellule: _Cell,
    texte: str,
    style: str,
    couleur: str | None = None,
    gras: bool | None = None,
    premier: bool = True,
) -> None:
    paragraphe = cellule.paragraphs[0] if premier else cellule.add_paragraph()
    paragraphe.style = style
    run = paragraphe.add_run(texte)
    _poser_police(run, style)
    if couleur:
        run.font.color.rgb = _rgb(couleur)
    if gras is not None:
        run.font.bold = gras


def saut_de_page(document: DocumentWord) -> None:
    """Saut explicite. La référence en compte 30 : le flux automatique ne suffit pas."""
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ── 1 et 7. Couverture et quatrième de couverture ────────────────────────────


def _fond_pleine_page(document: DocumentWord, couleur: str) -> None:
    """Rectangle pleine page ancré DERRIÈRE le texte, comme dans la référence."""
    run = document.add_paragraph().add_run()
    xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml">'
        f'<v:rect style="position:absolute;margin-left:0;margin-top:0;'
        f"width:{PAGE_LARGEUR_PT}pt;height:{PAGE_HAUTEUR_PT}pt;z-index:-251658752;"
        "mso-position-horizontal:left;mso-position-horizontal-relative:page;"
        'mso-position-vertical:top;mso-position-vertical-relative:page" '
        f'fillcolor="{couleur}" stroked="f"/></w:pict>'
    )
    run._r.append(parse_xml(xml))


def couverture(
    document: DocumentWord,
    palette: Palette,
    *,
    titre: str,
    sous_titre: str = "",
    client: str = "",
    mention: str = "Document confidentiel",
    logo: bytes | None = None,
) -> None:
    _fond_pleine_page(document, palette.primaire)

    for _ in range(4):
        document.add_paragraph()

    if logo:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(logo), width=Emu(1_600_000))

    p = document.add_paragraph(style=STYLE_TITRE_DOCUMENT)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titre.upper())
    _poser_police(run, STYLE_TITRE_DOCUMENT)
    run.font.color.rgb = _rgb(palette.texte_sur_primaire)

    if sous_titre:
        p = document.add_paragraph(style=STYLE_SOUS_TITRE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sous_titre)
        run.font.name = POLICE_CORPS
        run.font.color.rgb = _rgb(palette.rose_grise)

    if client:
        p = document.add_paragraph(style=STYLE_SOUS_TITRE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(client)
        run.font.name = POLICE_CORPS
        run.font.color.rgb = _rgb(palette.texte_sur_primaire)

    for _ in range(10):
        document.add_paragraph()

    p = document.add_paragraph(style=STYLE_LEGENDE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(mention)
    run.font.name = POLICE_CORPS
    run.font.color.rgb = _rgb(palette.rose_grise)

    saut_de_page(document)


def quatrieme_couverture(
    document: DocumentWord,
    palette: Palette,
    *,
    mentions: Sequence[str],
    logo: bytes | None = None,
) -> None:
    saut_de_page(document)
    _fond_pleine_page(document, palette.primaire)

    # Douze, comme relevé sur la référence. J'ai un temps ramené ce nombre à
    # huit en croyant que la quatrième de couverture débordait sur une page
    # vide. Elle ne débordait pas : ma mesure écartait toute ligne contenant
    # « confidentiel », pour retirer le pied de page — et effaçait donc le
    # contenu même de cette page, qui n'est fait que de ces mentions. La page
    # comptée vide était la quatrième de couverture faisant son travail.
    for _ in range(12):
        document.add_paragraph()

    if logo:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(logo), width=Emu(1_200_000))

    for mention in mentions:
        p = document.add_paragraph(style=STYLE_LEGENDE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(mention)
        run.font.name = POLICE_CORPS
        run.font.color.rgb = _rgb(palette.rose_grise)


# ── 2. Bandeau de chapitre ───────────────────────────────────────────────────


def marqueur_de_chapitre(numero: int) -> str:
    """Texte exact du bandeau qui ouvre un chapitre dans le `.docx`.

    Exporté, et c'est tout l'objet de cette fonction : la passe de vérification
    cherche ce marqueur dans le document livré pour dire si un chapitre est
    présent. Elle en portait sa PROPRE version, « Chapitre 01 » en casse mixte,
    quand le rendu écrit « CHAPITRE 01 » en capitales. La comparaison est
    sensible à la casse : le contrôle déclarait donc les vingt-trois chapitres
    ABSENTS d'un document qui les contient tous, en bloquant.

    Mesuré sur la fixture complète : 604 Ko, 128 tableaux, les vingt-trois
    bandeaux présents, et « Chapitre(s) absent(s) du livrable : [0, 1, ..., 22] ».
    Aucune étude de marché ne pouvait partir, et le motif était introuvable
    dans le document par qui l'ouvrait — c'est la règle 2 mot pour mot.

    Deux modules qui décrivent le même texte finissent toujours par diverger
    (règle 5). Il n'y en a plus qu'un : celui qui l'écrit.
    """
    return f"CHAPITRE {numero:02d}"


def bandeau_chapitre(
    document: DocumentWord, palette: Palette, numero: int, titre: str, accroche: str = ""
) -> None:
    """Tableau 1×1 pleine largeur, fond prune, texte blanc en capitales."""
    table = _table(document, 1, 1, [LARGEUR_UTILE_DXA])
    marges_cellules(table, haut=280, cote=280)
    cellule = table.rows[0].cells[0]
    fond_cellule(cellule, palette.fond_bandeau)

    _ecrire(
        cellule, marqueur_de_chapitre(numero), STYLE_ENCADRE_TITRE,
        couleur=palette.rose_grise,
    )
    _ecrire(
        cellule, titre.upper(), STYLE_BANDEAU,
        couleur=palette.texte_sur_primaire, premier=False,
    )
    if accroche:
        _ecrire(
            cellule, accroche, STYLE_ENCADRE_CORPS,
            couleur=palette.rose_grise, premier=False,
        )
    document.add_paragraph()


# ── 3. Encadré étiqueté ──────────────────────────────────────────────────────


def encadre(
    document: DocumentWord,
    palette: Palette,
    libelle: str,
    lignes: Sequence[str],
    *,
    verdict: bool = False,
) -> None:
    """Tableau 1×2 : étiquette prune étroite à gauche, corps crème à droite."""
    largeur_libelle = 1800
    table = _table(document, 1, 2, [largeur_libelle, LARGEUR_UTILE_DXA - largeur_libelle])
    marges_cellules(table, haut=180, cote=200)

    etiquette = table.rows[0].cells[0]
    fond_cellule(etiquette, palette.fond_bandeau)
    _ecrire(
        etiquette, libelle.upper(), STYLE_ENCADRE_TITRE,
        couleur=palette.texte_sur_primaire,
    )

    corps = table.rows[0].cells[1]
    fond_cellule(corps, palette.fond_clair_alt if verdict else palette.fond_clair)
    for index, ligne in enumerate(lignes):
        _ecrire(
            corps, ligne, STYLE_ENCADRE_CORPS,
            couleur=palette.prune_fonce, premier=index == 0,
        )
    document.add_paragraph()


# ── 4. Grille de chiffres clés ───────────────────────────────────────────────


def grille_chiffres(
    document: DocumentWord, palette: Palette, chiffres: Sequence[tuple[str, str, str]]
) -> None:
    """Tableau 1×3 par rangée, alternance crème / rose pâle / crème."""
    fonds = (palette.fond_clair, palette.rose_pale, palette.fond_clair)
    largeur = LARGEUR_UTILE_DXA // 3

    for depart in range(0, len(chiffres), 3):
        tranche = list(chiffres[depart : depart + 3])
        table = _table(document, 1, 3, [largeur] * 3)
        marges_cellules(table, haut=200, cote=160)
        for index in range(3):
            cellule = table.rows[0].cells[index]
            if index >= len(tranche):
                fond_cellule(cellule, "#FFFFFF")
                continue
            valeur, libelle, source = tranche[index]
            fond_cellule(cellule, fonds[index])
            _ecrire(cellule, valeur, STYLE_CHIFFRE_VALEUR, couleur=palette.primaire)
            _ecrire(
                cellule, libelle, STYLE_CHIFFRE_LIBELLE,
                couleur=palette.prune_fonce, premier=False,
            )
            if source:
                _ecrire(
                    cellule, source, STYLE_CHIFFRE_LIBELLE,
                    couleur=palette.texte_legende, premier=False,
                )
        document.add_paragraph()


# ── 5. Tableau de données ────────────────────────────────────────────────────


#: Part minimale et maximale d'une colonne, en fraction de la largeur utile.
#:
#: Le plancher empêche qu'une colonne « Oui / Non » soit réduite à un filet
#: illisible ; le plafond empêche qu'une colonne de commentaire mange tout et
#: écrase les autres. Entre les deux, la place suit le contenu.
_PART_MIN = 0.07
_PART_MAX = 0.34


def _largeurs(
    entetes: Sequence[str], lignes: Sequence[Sequence[str]], colonnes: int
) -> list[int]:
    """Largeur de chaque colonne, PROPORTIONNELLE à ce qu'elle porte.

    ## Le défaut corrigé

    La largeur était `LARGEUR_UTILE_DXA // colonnes` : parts strictement
    égales. Un tableau à neuf colonnes — le modèle de référence en porte un au
    chapitre 19 — donnait donc neuf colonnes d'un neuvième, quelle que soit leur
    matière. Une colonne « Priorité » contenant « Haute » recevait autant de
    place qu'une colonne « Recommandation » de deux lignes, qui se retrouvait
    coupée en accordéon.

    Signalé par la cliente le 09/08/2026 : « problème de tableaux trop étroits,
    pages 10-11 ».

    ## La règle

    Le poids d'une colonne est la longueur de son texte le plus long — en-tête
    compris, car un en-tête long élargit la colonne autant qu'une cellule. On
    borne ensuite chaque part entre un plancher et un plafond, puis on
    renormalise pour que la somme fasse exactement la largeur utile : sans cette
    dernière étape, les bornes feraient déborder ou rétrécir le tableau.
    """
    textes: list[int] = []
    for index in range(colonnes):
        candidats = [len(entetes[index]) if index < len(entetes) else 0]
        candidats += [
            len(ligne[index]) for ligne in lignes if index < len(ligne)
        ]
        # Une colonne vide vaut 1 et non 0 : à zéro partout, la division qui
        # suit n'aurait pas de sens.
        textes.append(max(max(candidats, default=1), 1))

    total = sum(textes)
    parts = [
        min(max(poids / total, _PART_MIN), _PART_MAX) for poids in textes
    ]
    somme = sum(parts)
    return [max(int(LARGEUR_UTILE_DXA * part / somme), 1) for part in parts]


def tableau(
    document: DocumentWord,
    palette: Palette,
    entetes: Sequence[str],
    lignes: Sequence[Sequence[str]],
    source: str = "",
) -> None:
    """En-tête prune sur texte blanc, corps crème, bordures fines."""
    colonnes = max(len(entetes), 1)
    table = _table(
        document, 1 + len(lignes), colonnes, _largeurs(entetes, lignes, colonnes)
    )
    _bordures(table, palette.fond_clair_alt, epaisseur=4)
    marges_cellules(table, haut=90, cote=140)
    # Seul le tableau de DONNÉES a une ligne d'en-tête à redonner : les autres
    # composants passant par `_table` (encadré, grille de chiffres) n'en ont pas.
    ne_pas_fractionner(table, entete_repetee=True)

    for index, intitule in enumerate(entetes):
        cellule = table.rows[0].cells[index]
        fond_cellule(cellule, palette.fond_bandeau)
        _ecrire(
            cellule, intitule, STYLE_TABLEAU_ENTETE,
            couleur=palette.texte_sur_primaire, gras=True,
        )

    for numero, ligne in enumerate(lignes):
        for index in range(colonnes):
            cellule = table.rows[numero + 1].cells[index]
            fond_cellule(cellule, palette.fond_clair)
            _ecrire(
                cellule,
                ligne[index] if index < len(ligne) else "",
                STYLE_TABLEAU_CELLULE,
                couleur=palette.texte_corps,
            )

    if source:
        note_source(document, palette, source)
    else:
        document.add_paragraph()


# ── 6. Graphique ─────────────────────────────────────────────────────────────


def graphique(
    document: DocumentWord, palette: Palette, png: bytes, titre: str = "", source: str = ""
) -> None:
    """Insère un PNG matplotlib sur toute la largeur utile."""
    if titre:
        # Une LÉGENDE, pas un en-tête de tableau — et pas un titre de section
        # non plus : le titre d'une figure n'a rien à faire dans la table des
        # matières, il nomme une image, pas une partie du document.
        p = document.add_paragraph(style=STYLE_LEGENDE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titre)
        run.font.name = POLICE_CORPS
        run.font.color.rgb = _rgb(palette.primaire)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.italic = False
        # Un titre de graphique séparé de son image est pire qu'un sous-titre
        # orphelin : il annonce une figure absente de la page.
        garder_avec_la_suite(p)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(png), width=LARGEUR_UTILE_EMU)
    if source:
        # L'image tient sa source avec elle ; sans quoi la légende part seule
        # en haut de la page suivante.
        garder_avec_la_suite(p)

    if source:
        note_source(document, palette, source)


# ── 8. Sommaire ──────────────────────────────────────────────────────────────


def sommaire(
    document: DocumentWord, palette: Palette, entrees: Sequence[tuple[str, str, str]]
) -> None:
    """Numéro et intitulé de chaque chapitre. La référence n'emploie pas de champ TOC.

    La colonne « Page » a été RETIRÉE. Elle était rendue à chaque étude et
    restait **toujours vide** : `depuis_json.rendre_etude` passe une chaîne
    vide en troisième position, et rien nulle part ne calcule une pagination —
    elle n'existe qu'après la mise en page, que ce moteur ne fait pas.

    Une colonne intitulée « Page » sans un seul numéro n'est pas un détail de
    forme : elle annonce au lecteur une information, puis ne la donne pas. Le
    tenir pour un défaut mineur, c'est le laisser dans le document livré.

    La navigation par chapitre existe désormais pour de bon, et par le moyen
    que Word attend : les styles de titre portent un niveau de plan (voir
    `gabarit._definir`), donc le volet de navigation les liste et une table des
    matières automatique se génère en deux clics, avec ses vraies pages.

    Le troisième élément de chaque entrée est conservé dans la signature : les
    appelants le fournissent déjà, et il portera la page le jour où la
    pagination sera connue.
    """
    p = document.add_paragraph(style=STYLE_BANDEAU)
    run = p.add_run("Sommaire")
    _poser_police(run, STYLE_BANDEAU)
    run.font.color.rgb = _rgb(palette.primaire)

    tableau(
        document, palette, ["Chap.", "Intitulé"],
        [[numero, titre] for numero, titre, _page in entrees],
    )


# ── Éléments de texte ────────────────────────────────────────────────────────


def sous_titre(document: DocumentWord, palette: Palette, texte: str) -> None:
    """Titre de section numéroté (« 1.1 Deux périmètres à ne pas confondre »).

    Il portait le style des EN-TÊTES DE TABLEAU, pendant que « Étude Titre
    section » — défini pour cet usage exact — n'était employé nulle part. Deux
    conséquences : le volet de navigation de Word restait vide, et le nom de
    style qu'affiche le ruban annonçait un tableau au lecteur qui cliquait dans
    un titre.
    """
    p = document.add_paragraph(style=STYLE_SECTION)
    run = p.add_run(texte)
    run.font.name = POLICE_CORPS
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = _rgb(palette.primaire)
    # Un sous-titre seul en bas de page est la faute de mise en page la plus
    # visible : le lecteur tourne la page pour trouver ce qu'il annonce.
    garder_avec_la_suite(p)


def paragraphe(document: DocumentWord, palette: Palette, texte: str) -> None:
    p = document.add_paragraph(style=STYLE_CORPS)
    run = p.add_run(texte)
    _poser_police(run, STYLE_CORPS)
    run.font.color.rgb = _rgb(palette.texte_corps)


#: Une URL complète dans le texte. Le motif est large à dessein : on veut
#: l'attraper qu'elle porte ou non son protocole.
_URL = re.compile(r"(?:https?://|www\.)[^\s,;)\]]+", re.IGNORECASE)

#: Sous-domaines qui ne disent rien du nom de la source.
_PREFIXES_INUTILES = ("www.", "fr.", "en.", "m.")


def source_lisible(texte: str) -> str:
    """Remplace une URL par le NOM de la source qu'elle désigne.

    « https://www.fevad.com/wp-content/uploads/2025/06/chiffres-cles-2025.pdf »
    devient « fevad.com ».

    ## Pourquoi

    Signalé par la cliente le 09/08/2026 : « à la page 40, les sources sont
    écrites sous forme d'URL longues, parfois avec des caractères parasites.
    Visuellement, il vaut mieux utiliser des noms de source propres. »

    Elle a raison, et pas seulement pour l'esthétique : une URL de cent
    cinquante signes déborde de la colonne, force un retour à la ligne au milieu
    d'un mot, et n'apprend rien au lecteur que le nom du site ne lui dise mieux.
    Personne ne recopie une URL depuis un PDF.

    ## Ce qu'on ne fait pas

    On ne supprime pas la source : « Fevad, 2025 » reste « Fevad, 2025 ». Seule
    l'adresse est remplacée par son domaine, et le reste de la phrase traverse
    intact. Retirer la source entière ferait perdre l'information au nom de la
    mise en page — l'inverse de ce qui est demandé.
    """
    def _domaine(correspondance: re.Match[str]) -> str:
        url = correspondance.group(0)
        hote = url.split("//", 1)[-1].split("/", 1)[0].lower()
        for prefixe in _PREFIXES_INUTILES:
            if hote.startswith(prefixe):
                hote = hote[len(prefixe):]
        return hote or url

    return _URL.sub(_domaine, texte)


def note_source(document: DocumentWord, palette: Palette, texte: str) -> None:
    p = document.add_paragraph(style=STYLE_SOURCE)
    run = p.add_run(source_lisible(texte))
    _poser_police(run, STYLE_SOURCE)
    run.font.color.rgb = _rgb(palette.texte_legende)


def legende(document: DocumentWord, palette: Palette, texte: str) -> None:
    p = document.add_paragraph(style=STYLE_LEGENDE)
    run = p.add_run(texte)
    _poser_police(run, STYLE_LEGENDE)
    run.font.color.rgb = _rgb(palette.texte_legende)


# ── 9. Matrice à quatre cases ────────────────────────────────────────────────


def matrice_quadrants(
    document: DocumentWord,
    palette: Palette,
    quadrants: Sequence[tuple[str, Sequence[str]]],
) -> None:
    """Grille 2×2 : un intitulé coloré par case, puis des lignes courtes.

    Volontairement générique plutôt que « SWOT » : la même forme sert au SWOT,
    à la matrice d'Ansoff, au couple probabilité/impact, à la grille
    effort/gain. Nommer le composant d'après l'un de ses usages aurait conduit
    à en écrire quatre (règle 4 : viser la classe du défaut).

    Attend exactement quatre cases ; au-delà, le surplus est ignoré, en deçà
    les cases manquantes restent blanches.
    """
    largeur = LARGEUR_UTILE_DXA // 2
    fonds = (
        palette.fond_clair,
        palette.rose_pale,
        palette.rose_pale,
        palette.fond_clair,
    )
    table = _table(document, 2, 2, [largeur, largeur])
    _bordures(table, palette.fond_clair_alt, epaisseur=8)
    marges_cellules(table, haut=160, cote=180)

    for index in range(4):
        cellule = table.rows[index // 2].cells[index % 2]
        if index >= len(quadrants):
            fond_cellule(cellule, "#FFFFFF")
            continue
        intitule, lignes = quadrants[index]
        fond_cellule(cellule, fonds[index])
        _ecrire(
            cellule, intitule.upper(), STYLE_ENCADRE_TITRE,
            couleur=palette.primaire, gras=True,
        )
        for ligne in lignes:
            _ecrire(
                cellule, ligne, STYLE_ENCADRE_CORPS,
                couleur=palette.prune_fonce, premier=False,
            )
    document.add_paragraph()


# ── 10. Barre de répartition ─────────────────────────────────────────────────


def barre_repartition(
    document: DocumentWord,
    palette: Palette,
    parts: Sequence[tuple[str, float]],
    source: str = "",
) -> None:
    """Bande horizontale découpée au prorata des parts, libellés en dessous.

    Lit une répartition d'un coup d'œil, sans occuper la place d'un graphique.
    Les largeurs sont calculées en DXA à partir des poids fournis, quel que
    soit leur total : des parts en euros comme en pourcentage donnent la même
    bande.

    Deux rangées, et non une seule : la bande colorée porte le pourcentage, la
    rangée du dessous porte les libellés sur fond blanc. C'est plus lisible
    qu'un libellé long dans une case étroite, et cela donne à ce composant une
    forme (2 rangées) que rien d'autre du livrable n'a — un encadré est une
    rangée de deux cellules, une grille de chiffres une rangée de trois. Sans
    cette distinction, toute mesure du document qui classe les tableaux par
    leur forme confondrait les trois.
    """
    poids = [max(valeur, 0.0) for _, valeur in parts]
    total = sum(poids) or 1.0
    largeurs = [max(int(LARGEUR_UTILE_DXA * valeur / total), 900) for valeur in poids]
    # Le minimum de lisibilité impose de rogner la case la plus large plutôt
    # que de laisser le tableau déborder de la largeur utile.
    ecart = sum(largeurs) - LARGEUR_UTILE_DXA
    if ecart > 0:
        largeurs[largeurs.index(max(largeurs))] -= ecart

    couleurs = palette.series_graphique
    table = _table(document, 2, len(largeurs), largeurs)
    marges_cellules(table, haut=120, cote=90)
    for index, (libelle, valeur) in enumerate(parts):
        fond = couleurs[index % len(couleurs)]
        bande = table.rows[0].cells[index]
        fond_cellule(bande, fond)
        _ecrire(
            bande, f"{valeur / total * 100:.0f} %", STYLE_CHIFFRE_LIBELLE,
            couleur=texte_lisible_sur(fond), gras=True,
        )
        etiquette = table.rows[1].cells[index]
        fond_cellule(etiquette, "#FFFFFF")
        _ecrire(etiquette, libelle, STYLE_LEGENDE, couleur=palette.texte_legende)

    if source:
        note_source(document, palette, source)
    else:
        document.add_paragraph()


def liste(document: DocumentWord, palette: Palette, elements: Sequence[str]) -> None:
    """Puces via la numérotation Word — jamais de caractère de puce en dur."""
    for element in elements:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(element)
        run.font.name = POLICE_CORPS
        run.font.color.rgb = _rgb(palette.texte_corps)


# ── 9. Business Model Canvas ─────────────────────────────────────────────────

#: La disposition d'Osterwalder, telle que l'AFE la distribue.
#:
#: Demande de la cliente le 13/08/2026, modèle à l'appui : « il serait plus
#: intéressant de mettre un business model canvas complété dans sa forme
#: originale qu'un tableau ».
#:
#: Elle a raison, et ce n'est pas une question d'esthétique. Le canvas se LIT
#: par blocs : ce que l'entreprise fait est à gauche, ce que le client reçoit
#: à droite, l'argent en bas. Neuf lignes empilées perdent cette lecture — on
#: obtient une liste, pas une carte, et le lecteur ne voit plus d'un coup d'œil
#: qu'une proposition de valeur sans canal pour l'atteindre est un modèle qui
#: ne tient pas.
#:
#: Grille : cinq colonnes, trois rangées. Les blocs pleine hauteur occupent
#: leurs deux premières rangées, les colonnes 1 et 3 se coupent en deux, et la
#: rangée du bas se partage entre coûts et revenus.
_CASES_DU_CANVAS: tuple[tuple[str, str, int, int, int, int], ...] = (
    # (champ, intitulé, ligne, colonne, hauteur, largeur)
    ("partenaires_cles",    "Partenaires clés",       0, 0, 2, 1),
    ("activites_cles",      "Activités clés",         0, 1, 1, 1),
    ("ressources_cles",     "Ressources clés",        1, 1, 1, 1),
    ("proposition_valeur",  "Proposition de valeur",  0, 2, 2, 1),
    ("relation_client",     "Relation client",        0, 3, 1, 1),
    ("canaux",              "Canaux",                 1, 3, 1, 1),
    ("segments_clientele",  "Segments de clientèle",  0, 4, 2, 1),
    ("structure_couts",     "Structure de coûts",     2, 0, 1, 2),
    ("sources_revenus",     "Sources de revenus",     2, 2, 1, 3),
)


def business_model_canvas(
    document: DocumentWord,
    palette: Palette,
    blocs: dict[str, Sequence[str]],
    source: str = "",
) -> None:
    """Dessine le canvas dans sa disposition d'origine, cases fusionnées.

    `blocs` associe chaque champ du canvas à ses éléments. Un bloc vide garde
    sa case : un canvas amputé d'une case cesse d'être un canvas, et le trou
    DIT quelque chose — un modèle sans partenaires clés est une information,
    pas une omission de mise en page.
    """
    largeur = LARGEUR_UTILE_DXA // 5
    table = _table(document, 3, 5, [largeur] * 5)
    _bordures(table, palette.fond_clair_alt, epaisseur=4)
    marges_cellules(table, haut=90, cote=120)

    for champ, intitule, ligne, colonne, hauteur, largeur_cases in _CASES_DU_CANVAS:
        cellule = table.rows[ligne].cells[colonne]
        if hauteur > 1 or largeur_cases > 1:
            cellule = cellule.merge(
                table.rows[ligne + hauteur - 1].cells[colonne + largeur_cases - 1]
            )
        fond_cellule(cellule, palette.fond_clair)

        # L'intitulé de la case, puis ses éléments en puces. On passe par
        # `_ecrire` plutôt que de poser le style à la main : lui seul connaît
        # la police du gabarit, et un style appliqué sans elle rend un texte
        # qui détonne au milieu du document.
        _ecrire(
            cellule, intitule, STYLE_TABLEAU_ENTETE,
            couleur=palette.primaire, gras=True,
        )
        for element in blocs.get(champ) or ():
            texte = str(element).strip()
            if texte:
                _ecrire(
                    cellule, f"• {texte}", STYLE_TABLEAU_CELLULE,
                    couleur=palette.texte_corps, premier=False,
                )

    if source:
        note_source(document, palette, source)
    else:
        document.add_paragraph()
