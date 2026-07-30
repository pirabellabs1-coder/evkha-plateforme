"""Construction et chargement du gabarit Word.

Le gabarit porte la STRUCTURE : mise en page, styles nommés, en-tête, pied de
page. Il ne porte aucune couleur client — celles-ci sont appliquées à la
génération, puisqu'elles changent à chaque étude.

Toutes les mesures viennent du document de référence
« Etude_de_marche_Joalie_EVKHA_2026_v4.docx », relevées dans son XML.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentWord
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

#: Racine des gabarits, hors du code Python — ils se modifient dans Word.
CHEMIN_GABARITS = Path(__file__).resolve().parents[3] / "gabarits"
GABARIT_DEFAUT = CHEMIN_GABARITS / "livrable_evkha.docx"

# ── Mesures relevées dans le document de référence ───────────────────────────
POLICE_CORPS = "Aptos"
POLICE_TITRES = "Georgia"

MARGE_HAUT = Cm(2.0)
MARGE_BAS = Cm(1.7)
MARGE_GAUCHE = Cm(2.0)
MARGE_DROITE = Cm(2.0)
DISTANCE_ENTETE = Cm(0.7)
DISTANCE_PIED = Cm(0.7)

#: Couleur neutre des styles du gabarit. Elle est REMPLACÉE à la génération.
#: Volontairement grise : si elle apparaissait telle quelle dans un livrable,
#: le défaut se verrait immédiatement au lieu de passer pour un choix.
NEUTRE = RGBColor(0x20, 0x20, 0x20)

#: Noms des styles que le moteur de rendu attend. Contrat entre le gabarit et
#: le code : renommer un style dans Word casse le rendu, d'où le contrôle
#: `styles_manquants()`.
STYLE_TITRE_DOCUMENT = "EVKHA Titre document"
STYLE_SOUS_TITRE = "EVKHA Sous-titre"
STYLE_CHAPITRE = "EVKHA Titre chapitre"
STYLE_SECTION = "EVKHA Titre section"
STYLE_CORPS = "EVKHA Corps"
STYLE_BANDEAU = "EVKHA Bandeau"
STYLE_ENCADRE_TITRE = "EVKHA Encadré titre"
STYLE_ENCADRE_CORPS = "EVKHA Encadré corps"
STYLE_CHIFFRE_VALEUR = "EVKHA Chiffre valeur"
STYLE_CHIFFRE_LIBELLE = "EVKHA Chiffre libellé"
STYLE_TABLEAU_ENTETE = "EVKHA Tableau en-tête"
STYLE_TABLEAU_CELLULE = "EVKHA Tableau cellule"
STYLE_SOURCE = "EVKHA Source"
STYLE_LEGENDE = "EVKHA Légende"

STYLES_ATTENDUS: tuple[str, ...] = (
    STYLE_TITRE_DOCUMENT, STYLE_SOUS_TITRE, STYLE_CHAPITRE, STYLE_SECTION,
    STYLE_CORPS, STYLE_BANDEAU, STYLE_ENCADRE_TITRE, STYLE_ENCADRE_CORPS,
    STYLE_CHIFFRE_VALEUR, STYLE_CHIFFRE_LIBELLE, STYLE_TABLEAU_ENTETE,
    STYLE_TABLEAU_CELLULE, STYLE_SOURCE, STYLE_LEGENDE,
)


class GabaritIntrouvableError(FileNotFoundError):
    """Le fichier de gabarit n'existe pas."""


class GabaritIncompletError(RuntimeError):
    """Le gabarit ne porte pas tous les styles attendus par le moteur."""


def _definir(
    document: DocumentWord,
    nom: str,
    *,
    police: str,
    taille: float,
    gras: bool = False,
    italique: bool = False,
    couleur: RGBColor = NEUTRE,
    avant: float = 0,
    apres: float = 0,
    interligne: float = 1.15,
    alignement: WD_ALIGN_PARAGRAPH | None = None,
    majuscules: bool = False,
    interlettrage: int = 0,
) -> None:
    """Crée un style de paragraphe, ou le remet à ces valeurs s'il existe."""
    styles = document.styles
    try:
        style = styles[nom]
    except KeyError:
        style = styles.add_style(nom, WD_STYLE_TYPE.PARAGRAPH)

    style.font.name = police
    style.font.size = Pt(taille)
    style.font.bold = gras
    style.font.italic = italique
    style.font.color.rgb = couleur

    # python-docx ne pose que la police latine ; sans ces deux attributs, Word
    # substitue une autre fonte pour les caractères accentués sur certaines
    # installations — défaut invisible en test, visible chez le client.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), police)
    rfonts.set(qn("w:cs"), police)

    if majuscules:
        caps = rpr.makeelement(qn("w:caps"), {})
        rpr.append(caps)
    if interlettrage:
        spacing = rpr.makeelement(qn("w:spacing"), {qn("w:val"): str(interlettrage)})
        rpr.append(spacing)

    paragraphe = style.paragraph_format
    paragraphe.space_before = Pt(avant)
    paragraphe.space_after = Pt(apres)
    paragraphe.line_spacing = interligne
    if alignement is not None:
        paragraphe.alignment = alignement


def _champ_page(paragraphe: Paragraph) -> None:
    """Insère le champ Word « numéro de page » (pas un nombre figé)."""
    run = paragraphe.add_run()
    debut = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instruction = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instruction.text = " PAGE "
    fin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._r.append(debut)
    run._r.append(instruction)
    run._r.append(fin)


def construire_gabarit(destination: Path | None = None) -> Path:
    """Écrit le gabarit et retourne son chemin."""
    destination = destination or GABARIT_DEFAUT
    destination.parent.mkdir(parents=True, exist_ok=True)

    document = Document()

    section = document.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin, section.bottom_margin = MARGE_HAUT, MARGE_BAS
    section.left_margin, section.right_margin = MARGE_GAUCHE, MARGE_DROITE
    section.header_distance, section.footer_distance = DISTANCE_ENTETE, DISTANCE_PIED

    normal = document.styles["Normal"]
    normal.font.name = POLICE_CORPS
    normal.font.size = Pt(10)
    normal.font.color.rgb = NEUTRE

    # ── Styles du contrat de rendu ───────────────────────────────────────────
    _definir(document, STYLE_TITRE_DOCUMENT, police=POLICE_TITRES, taille=30,
             apres=6, interligne=1.0)
    _definir(document, STYLE_SOUS_TITRE, police=POLICE_CORPS, taille=14,
             apres=18, interligne=1.2)
    _definir(document, STYLE_CHAPITRE, police=POLICE_TITRES, taille=18,
             avant=18, apres=8, interligne=1.1)
    _definir(document, STYLE_SECTION, police=POLICE_CORPS, taille=13, gras=True,
             avant=12, apres=4)
    _definir(document, STYLE_CORPS, police=POLICE_CORPS, taille=10,
             apres=6, interligne=1.25, alignement=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Bandeau de chapitre : texte clair sur aplat de couleur principale.
    _definir(document, STYLE_BANDEAU, police=POLICE_TITRES, taille=16,
             avant=4, apres=4, interligne=1.1)

    _definir(document, STYLE_ENCADRE_TITRE, police=POLICE_CORPS, taille=9,
             gras=True, apres=2, majuscules=True, interlettrage=20)
    _definir(document, STYLE_ENCADRE_CORPS, police=POLICE_CORPS, taille=9.5,
             apres=0, interligne=1.2)

    _definir(document, STYLE_CHIFFRE_VALEUR, police=POLICE_TITRES, taille=20,
             apres=0, interligne=1.0, alignement=WD_ALIGN_PARAGRAPH.CENTER)
    _definir(document, STYLE_CHIFFRE_LIBELLE, police=POLICE_CORPS, taille=8,
             apres=0, interligne=1.1, alignement=WD_ALIGN_PARAGRAPH.CENTER)

    _definir(document, STYLE_TABLEAU_ENTETE, police=POLICE_CORPS, taille=9,
             gras=True, apres=0, interligne=1.1)
    _definir(document, STYLE_TABLEAU_CELLULE, police=POLICE_CORPS, taille=9,
             apres=0, interligne=1.15)

    _definir(document, STYLE_SOURCE, police=POLICE_CORPS, taille=8,
             italique=True, avant=2, apres=8)
    _definir(document, STYLE_LEGENDE, police=POLICE_CORPS, taille=8,
             italique=True, apres=6, alignement=WD_ALIGN_PARAGRAPH.CENTER)

    # ── En-tête et pied de page ──────────────────────────────────────────────
    # Les textes sont des REPÈRES, remplacés à la génération : l'en-tête porte
    # le nom du client final, qui change à chaque étude.
    entete = section.header.paragraphs[0]
    entete.text = "{{ client }}  /  {{ titre_document }}"
    entete.style = document.styles["Header"]
    entete.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pied = section.footer.paragraphs[0]
    pied.text = "{{ mention_confidentialite }}  ·  "
    pied.style = document.styles["Footer"]
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _champ_page(pied)

    # Le corps reste vide : c'est le moteur qui l'écrit.
    for paragraphe in list(document.paragraphs):
        element = paragraphe._element
        element.getparent().remove(element)

    document.save(str(destination))
    return destination


def styles_manquants(document: DocumentWord) -> list[str]:
    """Styles attendus par le moteur et absents du gabarit.

    Contrôle explicite plutôt qu'un `KeyError` en pleine génération : un
    gabarit retouché dans Word peut avoir perdu un style, et il vaut mieux le
    savoir avant d'avoir payé vingt-deux appels au modèle.
    """
    presents = {style.name for style in document.styles}
    return [nom for nom in STYLES_ATTENDUS if nom not in presents]


def charger_gabarit(chemin: Path | None = None) -> DocumentWord:
    """Ouvre le gabarit et vérifie qu'il porte bien tous les styles attendus."""
    chemin = chemin or GABARIT_DEFAUT
    if not chemin.is_file():
        msg = (
            f"Gabarit introuvable : {chemin}. "
            "Le produire avec `python manage.py creer_gabarit`."
        )
        raise GabaritIntrouvableError(msg)

    document = Document(str(chemin))
    manquants = styles_manquants(document)
    if manquants:
        msg = (
            f"Le gabarit {chemin.name} ne porte pas les styles suivants, "
            f"attendus par le moteur de rendu : {', '.join(manquants)}."
        )
        raise GabaritIncompletError(msg)
    return document
