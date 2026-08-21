"""Produit le dossier Soleau au format Word A PARTIR DU DOSSIER HTML.

    .venv/Scripts/python.exe backend/scripts/schemas_soleau.py
    .venv/Scripts/python.exe backend/scripts/dossier_soleau.py

## Pourquoi une conversion, et non un second document

Les deux versions ont d'abord ete ECRITES SEPAREMENT. Chaque ajout n'allait
donc que dans l'une des deux, et elles ont diverge : le HTML portait le detail
des dix etapes et les explications des extraits de code, le Word portait le
detail des quatre passes et trois sections sur l'achat a l'unite. Deux
documents censes dire la meme chose, et qui ne la disaient pas.

`docs/dossier-soleau.html` est desormais la SOURCE UNIQUE. Ce script la lit et
en compose le Word. Modifier le dossier, c'est modifier le HTML ; le Word suit.

## La charte

Celle du site EVKHA — noir, jaune #F8C51C, blanc dominant — et non la charte
de reference des livrables clients. C'est la demande : le dossier doit
ressembler au site de la cliente.

Les schemas sont les MEMES que dans la version consultable : ils sont rendus
en image par `schemas_soleau.py`, qui les lit dans cette meme source.
"""
from __future__ import annotations

import os
import re
import sys
from pathlib import Path

RACINE = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(RACINE / "backend"))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "evkha.settings")

import django  # noqa: E402

django.setup()

from bs4 import BeautifulSoup, Tag  # noqa: E402
from docx.enum.text import WD_LINE_SPACING  # noqa: E402
from docx.shared import Pt  # noqa: E402

from generation.rendu_word import composants  # noqa: E402
from generation.rendu_word.depuis_json import _remplacer_reperes  # noqa: E402
from generation.rendu_word.gabarit import charger_gabarit  # noqa: E402
from generation.rendu_word.palette import construire_palette  # noqa: E402

SOURCE = RACINE / "docs" / "dossier-soleau.html"
SCHEMAS = RACINE / "docs" / "schemas"

#: Couleurs du site EVKHA, relevees sur evkha.fr.
SITE_NOIR = "#000000"
SITE_JAUNE = "#F8C51C"
SITE_CREME = "#FDF6DF"

document = charger_gabarit()
palette = construire_palette(
    primaire=SITE_NOIR, secondaire=SITE_JAUNE, fond_clair=SITE_CREME
)

_remplacer_reperes(document, {
    "{{ client }}": "EVKHA BUSINESS ET FORMATIONS",
    "{{ titre_document }}": "Dossier technique — dépôt Soleau",
    "{{ mention_confidentialite }}": "Document confidentiel",
})


def saut_de_page() -> None:
    """Un saut qui ne laisse pas de page blanche derriere lui.

    La rupture est posee AVANT le paragraphe. Un saut insere DANS le
    paragraphe le laisse commencer sur la page qu'on quitte : quand celle-ci
    est pleine, il deborde, occupe seul la page suivante, et renvoie le
    chapitre encore une page plus loin. Mesure avant correction : page 18 sur
    29, vide.
    """
    paragraphe = document.add_paragraph()
    format_ = paragraphe.paragraph_format
    format_.page_break_before = True
    format_.space_before = Pt(0)
    format_.space_after = Pt(0)
    format_.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    format_.line_spacing = Pt(1)
    paragraphe.add_run().font.size = Pt(1)


def code(extrait: str) -> None:
    """Un extrait de code : chasse fixe, interligne serre, leger retrait."""
    lignes = extrait.strip("\n").split("\n")
    for index, ligne in enumerate(lignes):
        paragraphe = document.add_paragraph()
        run = paragraphe.add_run(ligne or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        format_ = paragraphe.paragraph_format
        format_.space_after = Pt(0)
        format_.space_before = Pt(0)
        format_.left_indent = Pt(16)
        if index < len(lignes) - 1:
            composants.garder_avec_la_suite(paragraphe)
    document.add_paragraph()


def _texte(noeud: Tag | None) -> str:
    """Le texte d'un noeud, espaces normalises."""
    if noeud is None:
        return ""
    return re.sub(r"\s+", " ", noeud.get_text(" ", strip=True)).strip()


def _tableau(table: Tag) -> None:
    entetes = [_texte(c) for c in table.select("thead th")]
    lignes = [
        [_texte(c) for c in tr.find_all(["td", "th"])]
        for tr in table.select("tbody tr")
    ]
    if entetes and lignes:
        composants.tableau(document, palette, entetes, lignes)


def _definitions(dl: Tag, entetes: tuple[str, str]) -> None:
    """Une liste de definitions devient un tableau a deux colonnes."""
    lignes = [
        [_texte(bloc.find("dt")), _texte(bloc.find("dd"))]
        for bloc in dl.find_all("div", recursive=False)
        if bloc.find("dt") and bloc.find("dd")
    ]
    if lignes:
        composants.tableau(document, palette, list(entetes), lignes)


def _encadre(bloc: Tag) -> None:
    etiquette = bloc.find(class_="etiquette")
    libelle = _texte(etiquette) if etiquette else "À retenir"
    corps = [
        _texte(p) for p in bloc.find_all("p")
        if p is not etiquette and _texte(p)
    ]
    if corps:
        composants.encadre(
            document, palette, libelle, corps,
            verdict="clair" in (bloc.get("class") or []),
        )


def _figure(figure: Tag, numero: int) -> None:
    """Insere le schema deja rendu, avec sa legende."""
    image = SCHEMAS / f"schema-{numero:02d}.png"
    if not image.is_file():
        msg = (
            f"Le schema {numero:02d} n'a pas ete rendu. "
            "Lancer d'abord `backend/scripts/schemas_soleau.py`."
        )
        raise FileNotFoundError(msg)
    composants.graphique(
        document, palette, image.read_bytes(),
        source=_texte(figure.find("figcaption")),
    )


def _corps_de_section(section: Tag, compteur: dict[str, int]) -> None:
    for noeud in section.children:
        if not isinstance(noeud, Tag):
            continue
        # bs4 rend `class` en LISTE (attribut multi-valué). L'annotation le
        # dit : sans elle, mypy infère `Any` et le `in` ci-dessous ne vérifie
        # plus rien.
        brut: str | list[str] = noeud.get("class") or []
        classes: list[str] = brut if isinstance(brut, list) else [brut]
        # Le titre, le numero et l'accroche sont DEJA portes par le bandeau du
        # chapitre. Les reprendre ici les imprimait une seconde fois, trois
        # centimetres plus bas.
        if noeud.name == "h2" or "num" in classes or "intro" in classes:
            continue
        if noeud.name == "h3":
            composants.sous_titre(document, palette, _texte(noeud))
        elif noeud.name == "p":
            composants.paragraphe(document, palette, _texte(noeud))
        elif noeud.name == "pre":
            code(noeud.get_text())
        elif noeud.name == "figure":
            compteur["figure"] += 1
            _figure(noeud, compteur["figure"])
        elif noeud.name == "ul":
            composants.liste(
                document, palette, [_texte(li) for li in noeud.find_all("li")]
            )
        elif noeud.name == "dl":
            _definitions(noeud, ("Terme", "Définition"))
        elif noeud.name == "div" and ("encadre" in classes or "clair" in classes):
            _encadre(noeud)
        elif noeud.name == "div" and "cadre" in classes:
            table = noeud.find("table")
            if table is not None:
                _tableau(table)


def construire() -> None:
    soupe = BeautifulSoup(SOURCE.read_text(encoding="utf-8"), "html.parser")

    tete = soupe.find("header", class_="tete")
    if not isinstance(tete, Tag):
        raise SystemExit(
            f"{SOURCE.name} n'a pas de <header class=\"tete\"> : "
            "la couverture n'aurait pas de titre."
        )
    composants.couverture(
        document, palette,
        titre=_texte(tete.find("h1")),
        sous_titre="Dossier technique — architecture et développement",
        client="EVKHA BUSINESS ET FORMATIONS",
        mention="Pièce établie pour le dépôt d'une enveloppe Soleau auprès de l'INPI",
    )

    composants.sous_titre(document, palette, "Identification de la version déposée")
    identite = soupe.find("dl", class_="identite")
    if not isinstance(identite, Tag):
        raise SystemExit(
            f"{SOURCE.name} n'a pas de <dl class=\"identite\"> : le dossier "
            "partirait sans le commit ni l'empreinte de l'archive déposée."
        )
    _definitions(identite, ("Élément", "Valeur"))

    confidentialite = soupe.find("div", class_="encadre")
    if confidentialite is not None:
        _encadre(confidentialite)
    saut_de_page()

    sections = soupe.find_all("section")
    composants.sommaire(document, palette, [
        (_texte(s.find(class_="num")), _texte(s.find("h2")), "") for s in sections
    ])
    saut_de_page()

    compteur = {"figure": 0}
    for index, section in enumerate(sections):
        numero = re.sub(r"\D", "", _texte(section.find(class_="num")))
        intro = section.find("p", class_="intro")
        composants.bandeau_chapitre(
            document, palette,
            int(numero or index),
            _texte(section.find("h2")),
            _texte(intro),
        )
        document.add_paragraph()
        _corps_de_section(section, compteur)
        if index < len(sections) - 1:
            saut_de_page()

    composants.quatrieme_couverture(
        document, palette,
        mentions=[
            "Dossier technique établi le 21 août 2026.",
            "EVKHA BUSINESS ET FORMATIONS.",
            "Pièce destinée au dépôt d'une enveloppe Soleau auprès de l'INPI.",
            "Aucune clé d'API, aucun secret d'authentification et aucune donnée "
            "personnelle de client ne figure dans ce dossier ni dans l'archive "
            "qui l'accompagne.",
        ],
    )


if __name__ == "__main__":
    construire()
    sortie = RACINE.parent / "EVKHA-dossier-technique-Soleau.docx"
    document.save(str(sortie))
    print(f"Document ecrit : {sortie}")
